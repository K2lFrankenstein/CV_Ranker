
import os,fitz
import pandas as pd
import docx,ast,json
from openai import OpenAI
import configparser
from .models  import *
from django.contrib.auth.models import User
from CV_Ranker.settings import BASE_DIR

path_ads = os.path.join(BASE_DIR,"ETL",'some_values.ini')

config = configparser.ConfigParser()
config.read(path_ads)
dewq = config.get("Section1","key1")
client = OpenAI(
    api_key=dewq,
)









def remove_newlines(serie):
    serie = serie.replace('\n', ' ')
    serie = serie.replace('\\n', ' ')
    serie = serie.replace('  ', ' ')
    serie = serie.replace('  ', ' ')
    serie = serie.replace('  ', ' ')
    serie = serie.replace('  ', ' ')
    serie = serie.replace('  ', ' ')
    return serie

def EXTRACT_TEXT_FROM_PDF(filename):
   
    pdf_file =  fitz.open(filename)
    PDFFILEDATALIST = ""

    for Page_No,page in enumerate(pdf_file):
        pymupdf_text = ""
        pymupdf_text = pymupdf_text + page.get_text()
        ddt  = remove_newlines(pymupdf_text)
              
        PDFFILEDATALIST+=ddt

        
    return PDFFILEDATALIST

def EXTRACT_DATA_FROM_DOC(filename):
    doc = docx.Document(filename)
    table_data_list = []

    for p in doc.paragraphs:
        table_data_list.append(p.text)

    for i, table in enumerate(doc.tables, start=1):
        for y, row in enumerate(table.rows):
            row_data = [cell.text for cell in row.cells]
            rowdata =",".join(row_data)
            table_data_list.append(rowdata)


    
    return table_data_list

def CV_scorer(cv_data,jd_data):
   

    resume_Hilights = {"highlights":"3y experience in ML/experience | M.Tech/education | PyTorch/skill | award/publications | AWS/certification",
                       "score" : "<int>" 
                       }

    response = client.chat.completions.create(
            model="gpt-3.5-turbo-16k" ,
            
            messages=[
                {"role": "system", "content": "You are expert in matching resumes experience & skill keyword with job descriptions keyword and you return a Json object"},
                {"role": "user", "content": (
                    f"Given the following job description:\n{jd_data}\n"
                    f"and the CV:\n{cv_data}\n"
                    "Provide Just compatibility score out of 100"
                    f"and top-5 hilight of the CV in JUST 15 WORD only and in format {resume_Hilights}\n"
                    f"make sure that you provide JSON result in above given format. Strictly follow the format. Do not add any additional data/strings inside/outside the jSON"
                )}
            ],
        )

    response_message = response.choices[0].message.content.strip()
    print(response_message,type(response_message))

    res =   ast.literal_eval(response_message)       #response_message.split("\n")
    score =  res["score"]                           #int(res[0].split("=")[1])
    hilight = res["highlights"]                       #ast.literal_eval(res[1])
    print(score,hilight,type(res))

    return score,hilight

def REstructured(cv_data):
   

    restructure_CV = {"fullName":"string",
                      "contactInfo":"json object",
                      "education":"json object with number as keys",
                      "work experience":"json object with number as keys", 
                      "projects":"json object with number as keys",
                      "skills":"list",
                      "others":"json object with number as keys"},

    token_l =len(cv_data)

    response = client.chat.completions.create(
            model="gpt-4-turbo",
            response_format={ "type": "json_object" },
            messages=[
    {"role": "system", "content": "You are an expert in structuring resumes from raw data "},
                {"role": "user", "content": (
                    f"Given following raw data:\n{cv_data}\n can you recreate the resume in this format {restructure_CV}\n"
                    f"make sure that you provide result in above given formats."

                )}
            ],
            # max_tokens=token_l
        )

    response_message = response.choices[0].message.content
    # passerd_data = ast.literal_eval(response_message)

    print(response_message,type(response_message))


    return response_message


def file_handeller(list_data_file,path_to_cv,path_to_jd,email):

    return_answer = []

    temp = list_data_file[0]

    jd_f_name = os.path.join(path_to_jd,temp["filename"]+temp["extension"])

    if temp["extension"] == ".json":
        with open(jd_f_name, 'r') as fd:
                raw_data = json.load(fd) 
        jd_data = raw_data["Job Description"]
    else:
        jd_data = EXTRACT_TEXT_FROM_PDF(jd_f_name)

    jd_tittle = temp["Job Tittle"]

    f_key = User.objects.get(username = email)
    JD_obj = JobDescription.objects.create(jobTittle = jd_tittle, jobDetails = jd_data, UID = f_key)
    JD_obj.save()

    print("completed job tittle")

    for obj in list_data_file:

        if obj["is_cv"] == True or obj["is_cv"] == "true":
            cv_f_name = os.path.join(path_to_cv,obj["filename"]+obj["extension"])
            cv_data = EXTRACT_TEXT_FROM_PDF(cv_f_name)
        

            score,highlight = CV_scorer(cv_data,jd_data)

            restruct_cv = REstructured(cv_data)

            restruct_cv_obj = ast.literal_eval(restruct_cv)

            fullname= restruct_cv_obj["fullName"]
            skills = restruct_cv_obj["skills"]
            ct = restruct_cv_obj.get("contactInfo","NAN")
            work_exp = restruct_cv_obj.get("work experience","NAN")
            others = restruct_cv_obj.get("others","NAN")
            project = restruct_cv_obj.get("projects","NAN")
            education = restruct_cv_obj.get("education","NAN")

            

            # JOB DESCRIPTION

            # RESUME STRUCTURE:
            cv_obj = StructuredFormate.objects.create(
                UID = f_key,
                jobTittle = JD_obj,

                resumeName = fullname,
                resumeHighlights = highlight,
                resumeScore = score,
                isShared = False,

                skills = skills,
                contactInfo = ct,
                experience = work_exp,
                others = others,
                project = project,
                education = education
                )
            
            cv_obj.save()

            if cv_obj :
                print( "Success")
                return_answer.append(
                       cv_obj.id               
                )

            else:
                print( "Fail")

    return  return_answer          





def PUSH_Model(email,tittle):


    # ---------------------------------------- JD MODEL-----------------------------

    raw_data = EXTRACT_TEXT_FROM_PDF(r"C:\Users\k2lea\Downloads\JD_4.pdf")
    f_key = User.objects.get(username = email)
    obj = JobDescription.objects.create(jobTittle = tittle,jobDetails = raw_data,UID = f_key)
    obj.save()

    # ---------------------------------------- StructuredFormate MODEL-----------------------------


    ## this data needs to come from the API CALL
    score = 50 
    highlight = "3y experience in ML/experience | M.Tech/education | PyTorch/skill | award/publications | AWS/certification"
    
    with open(r"D:\psudo_desktop\RANDOM SEARCH\CV_Ranker\ETL\api_data.txt", 'r') as fd:
                    raw_data = fd.read() 
                
    data = ast.literal_eval(raw_data)

    fullname= data["fullName"]
    skills = data["skills"]
    ct = data.get("contactInfo","NAN")
    work_exp = data.get("work experience","NAN")
    others = data.get("others","NAN")
    project = data.get("projects","NAN")
    education = data.get("education","NAN")

    f_key = User.objects.get(username = email)
    Jf_key = JobDescription.objects.get()

    obj = StructuredFormate.objects.create(
        UID = f_key,
        jobTittle = Jf_key,

        resumeName = fullname,
        resumeHighlights = highlight,
        resumeScore = score,
        isShared = False,

        skills = skills,
        contactInfo = ct,
        experience = work_exp,
        others = others,
        project = project,
        education = education
    )
    obj.save()


    if obj:
        return "Success"
    else:
         return "Fail"

